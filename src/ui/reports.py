"""
Professional report generation UI for the Mining Maintenance System.

This screen reads the same user-managed data used across the app and exports
clean, professional reports to the app exports folder.
"""

from __future__ import annotations

import json
import logging
import os
import re
import subprocess
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any, Dict, List, Optional

import customtkinter as ctk
import pandas as pd
from tkinter import messagebox

try:
    from ..app_paths import data_dir, exports_dir
    from ..machine_store import evaluate_machine_status, load_machines, machine_history
    from ..mine_store import get_active_mine
except Exception:
    from app_paths import data_dir, exports_dir
    from machine_store import evaluate_machine_status, load_machines, machine_history
    from mine_store import get_active_mine

from . import theme as theme_mod
from .gradient import GradientPanel


logger = logging.getLogger(__name__)
DATA_DIR = data_dir()

REPORT_CHOICES = {
    "Combined Operations Report": "all",
    "Machine Register": "machines",
    "Maintenance Summary": "maintenance",
    "Maintenance Completion Log": "completions",
    "Operators Report": "operators",
    "Schedules Report": "schedules",
    "Users Report": "users",
}

FORMAT_CHOICES = {
    "Printable HTML": "html",
    "Excel Workbook": "xlsx",
    "CSV Export": "csv",
    "PDF Document": "pdf",
    "Word Document": "docx",
}

REPORT_PRESETS = [
    "Custom",
    "This Month",
    "Overdue Only",
    "Completion Log",
]


class ReportsFrame(ctk.CTkFrame):
    def __init__(self, parent):
        super().__init__(parent, fg_color=theme_mod.SIMPLE_PALETTE.get("card", "transparent"))
        self.parent = parent
        self._surface = theme_mod.SIMPLE_PALETTE.get("card", "#111827")
        self._surface_alt = "#0b1220"
        self._text_primary = "#e2e8f0"
        self._text_muted = "#94a3b8"
        self._accent = theme_mod.SIMPLE_PALETTE.get("primary", "#2563eb")
        self._summary_labels: Dict[str, ctk.CTkLabel] = {}
        self._last_generated_report: Optional[Path] = None
        self.report_type_label = "Maintenance Summary"
        self.export_format_label = "Printable HTML"
        self.report_scope = "all"
        self._build_ui()
        self._refresh_overview()

    def _card(self, parent=None, *, corner_radius: int = 16):
        return ctk.CTkFrame(parent or self, fg_color=self._surface, corner_radius=corner_radius)

    def _build_ui(self):
        header = GradientPanel(
            self,
            colors=getattr(theme_mod, "SECTION_GRADIENTS", {}).get("reports", ("#0f172a", "#1d4ed8", "#0891b2")),
            corner_radius=16,
            border_color="#1d2a3f",
        )
        header.pack(fill="x", padx=18, pady=(18, 12))
        ctk.CTkLabel(
            header.content,
            text="Reports & Documents",
            font=("Segoe UI Semibold", 24),
            text_color=self._text_primary,
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            header.content,
            text="Generate clean professional reports directly from the same live user data stored inside the app.",
            font=("Segoe UI", 13),
            text_color="#dbeafe",
        ).pack(anchor="w", padx=16, pady=(0, 12))

        stats = self._card()
        stats.pack(fill="x", padx=18, pady=(0, 12))
        stats_row = ctk.CTkFrame(stats, fg_color="transparent")
        stats_row.pack(fill="x", padx=12, pady=12)
        for idx, label in enumerate(("Machines", "Due / Overdue", "Operators", "Schedules")):
            card = ctk.CTkFrame(stats_row, fg_color=self._surface_alt, corner_radius=14)
            card.grid(row=0, column=idx, sticky="nsew", padx=6)
            stats_row.grid_columnconfigure(idx, weight=1)
            ctk.CTkLabel(card, text=label, font=("Segoe UI", 12), text_color=self._text_muted).pack(anchor="w", padx=14, pady=(12, 4))
            value = ctk.CTkLabel(card, text="0", font=("Segoe UI Semibold", 24), text_color=self._text_primary)
            value.pack(anchor="w", padx=14, pady=(0, 12))
            self._summary_labels[label] = value

        options = self._card()
        options.pack(fill="x", padx=18, pady=(0, 12))
        grid = ctk.CTkFrame(options, fg_color="transparent")
        grid.pack(fill="x", padx=12, pady=12)
        grid.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(grid, text="Report Type", font=("Segoe UI Semibold", 14), text_color=self._text_primary).grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.report_type_menu = ctk.CTkOptionMenu(
            grid,
            values=list(REPORT_CHOICES.keys()),
            command=self._on_report_option_change,
            height=38,
            font=("Segoe UI", 13),
            dropdown_font=("Segoe UI", 13),
            fg_color="#1d4ed8",
            button_color="#1e40af",
            button_hover_color="#1e3a8a",
        )
        self.report_type_menu.grid(row=0, column=1, sticky="ew", padx=(12, 0), pady=(0, 8))
        self.report_type_menu.set(self.report_type_label)

        ctk.CTkLabel(grid, text="Export Format", font=("Segoe UI Semibold", 14), text_color=self._text_primary).grid(row=1, column=0, sticky="w")
        self.format_menu = ctk.CTkOptionMenu(
            grid,
            values=list(FORMAT_CHOICES.keys()),
            command=self._on_format_option_change,
            height=38,
            font=("Segoe UI", 13),
            dropdown_font=("Segoe UI", 13),
            fg_color="#0f766e",
            button_color="#115e59",
            button_hover_color="#134e4a",
        )
        self.format_menu.grid(row=1, column=1, sticky="ew", padx=(12, 0))
        self.format_menu.set("Printable HTML")

        ctk.CTkLabel(grid, text="Preset", font=("Segoe UI Semibold", 14), text_color=self._text_primary).grid(row=2, column=0, sticky="w", pady=(10, 0))
        self.preset_menu = ctk.CTkOptionMenu(
            grid,
            values=REPORT_PRESETS,
            command=self._apply_preset,
            height=38,
            font=("Segoe UI", 13),
            dropdown_font=("Segoe UI", 13),
            fg_color="#7c3aed",
            button_color="#6d28d9",
            button_hover_color="#5b21b6",
        )
        self.preset_menu.grid(row=2, column=1, sticky="ew", padx=(12, 0), pady=(10, 0))
        self.preset_menu.set("Custom")

        ctk.CTkLabel(grid, text="Date Range", font=("Segoe UI Semibold", 14), text_color=self._text_primary).grid(row=3, column=0, sticky="w", pady=(10, 0))
        range_wrap = ctk.CTkFrame(grid, fg_color="transparent")
        range_wrap.grid(row=3, column=1, sticky="ew", padx=(12, 0), pady=(10, 0))
        range_wrap.grid_columnconfigure(1, weight=1)
        range_wrap.grid_columnconfigure(3, weight=1)

        self.date_from_var = ctk.StringVar(value="")
        self.date_to_var = ctk.StringVar(value="")
        ctk.CTkLabel(range_wrap, text="From", font=("Segoe UI", 12), text_color=self._text_muted).grid(row=0, column=0, sticky="w", padx=(0, 8))
        from_entry = ctk.CTkEntry(range_wrap, textvariable=self.date_from_var, placeholder_text="YYYY-MM-DD", height=36, font=("Segoe UI", 13))
        from_entry.grid(row=0, column=1, sticky="ew", padx=(0, 12))
        from_entry.bind("<KeyRelease>", lambda _event: self._set_custom_and_refresh())
        ctk.CTkLabel(range_wrap, text="To", font=("Segoe UI", 12), text_color=self._text_muted).grid(row=0, column=2, sticky="w", padx=(0, 8))
        to_entry = ctk.CTkEntry(range_wrap, textvariable=self.date_to_var, placeholder_text="YYYY-MM-DD", height=36, font=("Segoe UI", 13))
        to_entry.grid(row=0, column=3, sticky="ew")
        to_entry.bind("<KeyRelease>", lambda _event: self._set_custom_and_refresh())

        preview = self._card()
        preview.pack(fill="both", expand=True, padx=18, pady=(0, 12))
        top = ctk.CTkFrame(preview, fg_color="transparent")
        top.pack(fill="x", padx=12, pady=(12, 6))
        ctk.CTkLabel(top, text="Report Snapshot", font=("Segoe UI Semibold", 16), text_color=self._text_primary).pack(side="left")
        refresh_btn = ctk.CTkButton(top, text="Refresh Data", width=110, height=32, command=self._refresh_overview)
        refresh_btn.pack(side="right")

        self.preview_box = ctk.CTkTextbox(
            preview,
            height=260,
            fg_color=self._surface_alt,
            corner_radius=14,
            text_color=self._text_primary,
            font=("Consolas", 12),
            border_width=1,
            border_color="#1f2937",
        )
        self.preview_box.pack(fill="both", expand=True, padx=12, pady=(0, 12))

        footer = ctk.CTkFrame(self, fg_color="transparent")
        footer.pack(fill="x", padx=18, pady=(0, 18))
        self.status_label = ctk.CTkLabel(footer, text="", font=("Segoe UI", 12), text_color=self._text_muted)
        self.status_label.pack(side="left")
        open_last_btn = ctk.CTkButton(
            footer,
            text="Open Last Report",
            command=self._open_last_report,
            font=("Segoe UI Semibold", 13),
            height=40,
            fg_color="#1e40af",
            hover_color="#1d4ed8",
        )
        open_last_btn.pack(side="right", padx=(0, 10))
        open_exports_btn = ctk.CTkButton(
            footer,
            text="Open Exports Folder",
            command=self._open_exports_folder,
            font=("Segoe UI Semibold", 13),
            height=40,
            fg_color="#334155",
            hover_color="#475569",
        )
        open_exports_btn.pack(side="right", padx=(0, 10))
        generate_btn = ctk.CTkButton(
            footer,
            text="Generate Report",
            command=self.generate_report,
            font=("Segoe UI Semibold", 14),
            height=40,
            fg_color=self._accent,
            hover_color="#1d4ed8",
        )
        generate_btn.pack(side="right")

    def _on_report_option_change(self, value: str):
        self.report_type_label = value
        self._set_custom_and_refresh()

    def _on_format_option_change(self, value: str):
        self.export_format_label = value
        self._refresh_overview()

    def _set_custom_and_refresh(self):
        self.report_scope = "all"
        try:
            self.preset_menu.set("Custom")
        except Exception:
            pass
        self._refresh_overview()

    def _apply_preset(self, value: str):
        today = datetime.now()
        preset = str(value or "Custom")
        self.report_scope = "all"
        if preset == "This Month":
            start = today.replace(day=1)
            self.report_type_label = "Maintenance Summary"
            self.report_type_menu.set(self.report_type_label)
            self.date_from_var.set(start.strftime("%Y-%m-%d"))
            self.date_to_var.set(today.strftime("%Y-%m-%d"))
        elif preset == "Overdue Only":
            self.report_type_label = "Maintenance Summary"
            self.report_type_menu.set(self.report_type_label)
            self.date_from_var.set("")
            self.date_to_var.set("")
            self.report_scope = "overdue"
        elif preset == "Completion Log":
            self.report_type_label = "Maintenance Completion Log"
            self.report_type_menu.set(self.report_type_label)
            self.date_from_var.set("")
            self.date_to_var.set("")
        else:
            self.report_scope = "all"
        self._refresh_overview()

    def _load_json_records(self, filename: str) -> List[Dict[str, Any]]:
        path = DATA_DIR / filename
        if not path.exists():
            return []
        try:
            try:
                raw_text = path.read_text(encoding="utf-8")
                payload = json.loads(raw_text) or []
            except json.JSONDecodeError:
                # Some local exports may include BOM; accept those gracefully.
                payload = json.loads(path.read_text(encoding="utf-8-sig")) or []
            if isinstance(payload, list):
                return [item for item in payload if isinstance(item, dict)]
            return []
        except Exception:
            logger.exception("Failed to read %s", path)
            return []

    def _active_mine_context(self) -> Dict[str, str]:
        try:
            mine = get_active_mine() or {}
        except Exception:
            mine = {}
        mine_name = str(mine.get("mine_name") or "").strip()
        company_name = str(mine.get("company_name") or "").strip()
        quarry_type = str(mine.get("quarry_type") or "").strip()
        lease_area = str(mine.get("lease_area") or "").strip()
        address = str(mine.get("address") or "").strip()
        maps_link = str(mine.get("google_maps_link") or "").strip()
        display_name = mine_name or "Unassigned Mine"
        return {
            "mine_name": mine_name,
            "company_name": company_name,
            "quarry_type": quarry_type,
            "lease_area": lease_area,
            "address": address,
            "google_maps_link": maps_link,
            "display_name": display_name,
        }

    def _base_summary(self, now: datetime, machine_rows: List[Dict[str, Any]], due_rows: List[Dict[str, Any]], completion_rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        mine = self._active_mine_context()
        summary = [
            {"label": "Generated", "value": now.strftime("%Y-%m-%d %H:%M:%S")},
            {"label": "Mine", "value": mine["display_name"]},
            {"label": "Machines", "value": len(machine_rows)},
            {"label": "Due / Overdue", "value": len(due_rows)},
            {"label": "Completion Logs", "value": len(completion_rows)},
        ]
        if mine.get("company_name"):
            summary.append({"label": "Company", "value": mine["company_name"]})
        return summary

    def _load_users(self) -> List[Dict[str, Any]]:
        xlsx = DATA_DIR / "users.xlsx"
        if xlsx.exists():
            try:
                return pd.read_excel(xlsx).fillna("").to_dict("records")
            except Exception:
                logger.exception("Failed to read users.xlsx")
        return self._load_json_records("users.json")

    def _parse_date_value(self, value: Any) -> datetime | None:
        raw = str(value or "").strip()
        if not raw:
            return None
        for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%Y/%m/%d", "%d-%m-%Y", "%d/%m/%Y"):
            try:
                return datetime.strptime(raw, fmt)
            except Exception:
                continue
        try:
            return datetime.fromisoformat(raw.replace(" ", "T"))
        except Exception:
            return None

    def _active_date_range(self) -> tuple[datetime | None, datetime | None]:
        start = self._parse_date_value(self.date_from_var.get())
        end = self._parse_date_value(self.date_to_var.get())
        if end is not None and len(str(self.date_to_var.get() or "").strip()) == 10:
            end = end.replace(hour=23, minute=59, second=59)
        return start, end

    def _filter_rows_by_date(self, rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        start, end = self._active_date_range()
        if rows is None or (start is None and end is None):
            return rows or []

        date_candidates = (
            "Completed At",
            "Recorded At",
            "Due Date",
            "Generated",
            "created_at",
            "completed_at",
            "date",
            "due_at",
            "due_date",
            "next_maintenance",
        )

        filtered: List[Dict[str, Any]] = []
        for row in rows:
            row_dt = None
            for key in date_candidates:
                if key in row:
                    row_dt = self._parse_date_value(row.get(key))
                    if row_dt is not None:
                        break
            if row_dt is None:
                filtered.append(row)
                continue
            if start is not None and row_dt < start:
                continue
            if end is not None and row_dt > end:
                continue
            filtered.append(row)
        return filtered

    def _machine_rows(self) -> List[Dict[str, Any]]:
        rows = []
        for machine in load_machines():
            status_ctx = evaluate_machine_status(machine)
            rows.append(
                {
                    "Machine ID": machine.get("id") or "",
                    "Name": machine.get("name") or machine.get("model") or "",
                    "Type": machine.get("type") or "",
                    "Model": machine.get("model") or "",
                    "Company": machine.get("company") or "",
                    "Health": str(status_ctx.get("status") or "normal").title(),
                    "Trigger": str(status_ctx.get("trigger") or "-").title(),
                    "Running Hours": machine.get("current_hours") or machine.get("hours") or "",
                    "Next Due Hours": machine.get("next_due_hours") or "",
                    "Due Date": status_ctx.get("due_date") or machine.get("next_maintenance") or machine.get("due_date") or "",
                    "Operator Phone": machine.get("operator_phone") or "",
                }
            )
        return rows

    def _completion_rows(self, machines: List[Dict[str, Any]] | None = None) -> List[Dict[str, Any]]:
        machines = machines if machines is not None else load_machines()
        rows: List[Dict[str, Any]] = []
        for machine in machines:
            for item in machine_history(machine):
                rows.append(
                    {
                        "Completed At": str(item.get("completed_at") or "").replace("T", " "),
                        "Machine ID": item.get("machine_id") or machine.get("id") or "",
                        "Machine": item.get("machine_name") or machine.get("name") or machine.get("model") or "",
                        "Completed By": item.get("completed_by") or "",
                        "Previous Status": str(item.get("previous_status") or "").title(),
                        "Runtime At Service": item.get("current_hours") or "",
                        "Rolled Due Date": item.get("rolled_due_date") or "",
                        "Rolled Due Hours": item.get("rolled_next_due_hours") or "",
                        "Notes": item.get("completion_notes") or "",
                    }
                )
        rows.sort(key=lambda item: str(item.get("Completed At") or ""), reverse=True)
        return rows

    def _maintenance_due_rows(self, machines: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        rows = []
        for machine in machines:
            ctx = evaluate_machine_status(machine)
            status = str(ctx.get("status") or "normal").lower()
            if status not in {"maintenance", "due", "overdue", "critical"}:
                continue
            rows.append(
                {
                    "Machine ID": machine.get("id") or "",
                    "Machine": machine.get("name") or machine.get("model") or "",
                    "Health": status.title(),
                    "Trigger": str(ctx.get("trigger") or "-").title(),
                    "Due Date": ctx.get("due_date") or machine.get("next_maintenance") or machine.get("due_date") or "",
                    "Current Hours": ctx.get("current_hours") or machine.get("current_hours") or machine.get("hours") or "",
                    "Next Due Hours": ctx.get("next_due_hours") or machine.get("next_due_hours") or "",
                    "Operator Phone": machine.get("operator_phone") or "",
                }
            )
        if self.report_scope == "overdue":
            rows = [row for row in rows if str(row.get("Health") or "").lower() == "overdue"]
        status_rank = {"critical": 0, "overdue": 1, "due": 2, "maintenance": 3}
        rows.sort(key=lambda item: (status_rank.get(str(item.get("Health") or "").lower(), 9), str(item.get("Machine ID") or "")))
        return rows

    def _hour_entry_rows(self, limit: int = 25) -> List[Dict[str, Any]]:
        rows = self._load_json_records("hour_entries.json")
        rows.sort(
            key=lambda item: str(
                item.get("timestamp")
                or item.get("created_at")
                or item.get("date")
                or item.get("time")
                or ""
            ),
            reverse=True,
        )
        shaped = []
        for item in rows[:limit]:
            shaped.append(
                {
                    "Recorded At": item.get("timestamp") or item.get("created_at") or item.get("date") or "",
                    "Machine ID": item.get("machine_id") or item.get("machine") or "",
                    "Shift": item.get("shift") or "",
                    "Hour Reading": item.get("hour_reading") or item.get("hours") or item.get("running_hours") or "",
                    "Operator": item.get("operator") or item.get("operator_name") or "",
                }
            )
        return shaped

    def _build_report_bundle(self, report_key: str) -> Dict[str, Any]:
        machines = load_machines()
        operators = self._load_json_records("operators.json")
        schedules = self._load_json_records("schedules.json")
        users = self._load_users()
        hour_entries = self._hour_entry_rows()
        maintenance_tasks = self._load_json_records("maintenance_tasks.json")
        completion_rows = self._completion_rows(machines)
        due_rows = self._maintenance_due_rows(machines)
        machine_rows = self._machine_rows()
        now = datetime.now()
        mine = self._active_mine_context()

        summary = self._base_summary(now, machine_rows, due_rows, completion_rows)
        site_row: Dict[str, Any] = {}
        if mine.get("mine_name"):
            site_row["Active Mine"] = mine["mine_name"]
        if mine.get("company_name"):
            site_row["Company"] = mine["company_name"]
        if mine.get("quarry_type"):
            site_row["Quarry Type"] = mine["quarry_type"]
        if mine.get("lease_area"):
            site_row["Lease Area"] = mine["lease_area"]
        if mine.get("address"):
            site_row["Address"] = mine["address"]
        if mine.get("google_maps_link"):
            site_row["Google Maps"] = mine["google_maps_link"]
        site_section = [{
            "title": "Mine Site Context",
            "description": "Active site profile attached to this report and used across the app workspace.",
            "rows": [site_row] if site_row else [],
        }]

        if report_key == "machines":
            return {
                "title": "Machine Register",
                "subtitle": "User-managed machine asset register with live health status.",
                "generated_at": now,
                "summary": summary,
                "mine_context": mine,
                "sections": site_section + [
                    {"title": "Machine Register", "description": "All machines currently stored in the application.", "rows": self._filter_rows_by_date(machine_rows)},
                ],
            }

        if report_key == "operators":
            return {
                "title": "Operators Report",
                "subtitle": "Current operator contact records used by the maintenance workflow.",
                "generated_at": now,
                "summary": summary + [{"label": "Operators", "value": len(operators)}],
                "mine_context": mine,
                "sections": site_section + [
                    {"title": "Operators", "description": "User-managed operator directory.", "rows": self._filter_rows_by_date(operators)},
                ],
            }

        if report_key == "schedules":
            return {
                "title": "Schedules Report",
                "subtitle": "Saved schedule configuration and routing details from the app.",
                "generated_at": now,
                "summary": summary + [{"label": "Schedules", "value": len(schedules)}],
                "mine_context": mine,
                "sections": site_section + [
                    {"title": "Schedules", "description": "All saved schedules from the application store.", "rows": self._filter_rows_by_date(schedules)},
                ],
            }

        if report_key == "users":
            return {
                "title": "Users Report",
                "subtitle": "Application user records and role-ready identity list.",
                "generated_at": now,
                "summary": summary + [{"label": "Users", "value": len(users)}],
                "mine_context": mine,
                "sections": site_section + [
                    {"title": "Users", "description": "Current users available in the local application store.", "rows": self._filter_rows_by_date(users)},
                ],
            }

        if report_key == "completions":
            return {
                "title": "Maintenance Completion Log",
                "subtitle": "Printable history of completed maintenance activities per machine.",
                "generated_at": now,
                "summary": summary,
                "mine_context": mine,
                "sections": site_section + [
                    {"title": "Completion History", "description": "Saved maintenance completion records captured when work is closed.", "rows": self._filter_rows_by_date(completion_rows)},
                ],
            }

        if report_key == "maintenance":
            return {
                "title": "Maintenance Summary",
                "subtitle": "Live maintenance position combining date-based and runtime-based rules.",
                "generated_at": now,
                "summary": summary,
                "mine_context": mine,
                "sections": site_section + [
                    {"title": "Maintenance Priority Queue", "description": "Machines currently in maintenance, due, overdue, or critical states.", "rows": self._filter_rows_by_date(due_rows)},
                    {"title": "Recent Hour Entries", "description": "Latest runtime submissions from the worksheet flow.", "rows": self._filter_rows_by_date(hour_entries)},
                    {"title": "Open Maintenance Tasks", "description": "Tasks currently stored in the maintenance task file.", "rows": self._filter_rows_by_date(maintenance_tasks)},
                    {"title": "Completion History", "description": "Recent closed maintenance events captured from the machine workflow.", "rows": self._filter_rows_by_date(completion_rows[:50])},
                ],
            }

        return {
            "title": "Combined Operations Report",
            "subtitle": "Cross-functional export of machines, operators, schedules, users, hour entries, and completion history.",
            "generated_at": now,
            "summary": summary,
            "mine_context": mine,
            "sections": site_section + [
                {"title": "Machine Register", "description": "Current machine master data and live maintenance health.", "rows": self._filter_rows_by_date(machine_rows)},
                {"title": "Operators", "description": "Current operator directory.", "rows": self._filter_rows_by_date(operators)},
                {"title": "Schedules", "description": "Saved schedule records.", "rows": self._filter_rows_by_date(schedules)},
                {"title": "Users", "description": "Application user records.", "rows": self._filter_rows_by_date(users)},
                {"title": "Recent Hour Entries", "description": "Latest worksheet runtime entries.", "rows": self._filter_rows_by_date(hour_entries)},
                {"title": "Maintenance Completion History", "description": "Closed maintenance history captured from the machine screen.", "rows": self._filter_rows_by_date(completion_rows)},
            ],
        }

    def _refresh_overview(self):
        report_key = REPORT_CHOICES.get(self.report_type_label, "all")
        bundle = self._build_report_bundle(report_key)
        machines = load_machines()
        due_rows = self._filter_rows_by_date(self._maintenance_due_rows(machines))
        operators = self._filter_rows_by_date(self._load_json_records("operators.json"))
        schedules = self._filter_rows_by_date(self._load_json_records("schedules.json"))

        values = {
            "Machines": str(len(self._filter_rows_by_date(self._machine_rows()))),
            "Due / Overdue": str(len(due_rows)),
            "Operators": str(len(operators)),
            "Schedules": str(len(schedules)),
        }
        for key, label in self._summary_labels.items():
            try:
                label.configure(text=values.get(key, "0"))
            except Exception:
                pass

        preview_lines = [
            bundle["title"],
            bundle["subtitle"],
            "",
            f"Active mine: {bundle.get('mine_context', {}).get('display_name', 'Unassigned Mine')}",
            f"Preset: {self.preset_menu.get() if hasattr(self, 'preset_menu') else 'Custom'}",
            f"Export format: {self.export_format_label}",
            f"Date filter: {self.date_from_var.get().strip() or 'Any'} to {self.date_to_var.get().strip() or 'Any'}",
            f"Maintenance scope: {self.report_scope.title()}",
            f"Generated from app data folder: {DATA_DIR}",
            "",
            "Sections included:",
        ]
        for section in bundle.get("sections", []):
            preview_lines.append(f"  - {section['title']}: {len(section.get('rows', []))} rows")
            if section.get("description"):
                preview_lines.append(f"    {section['description']}")

        preview_lines.extend(["", "Summary:"])
        for item in bundle.get("summary", []):
            preview_lines.append(f"  - {item['label']}: {item['value']}")

        try:
            self.preview_box.delete("1.0", "end")
            self.preview_box.insert("1.0", "\n".join(preview_lines))
        except Exception:
            pass

    def generate_report(self):
        report_key = REPORT_CHOICES.get(self.report_type_label, "all")
        export_format = FORMAT_CHOICES.get(self.export_format_label, "html")
        bundle = self._build_report_bundle(report_key)
        timestamp = bundle["generated_at"].strftime("%Y%m%d_%H%M%S")
        site_name = re.sub(r"[^a-z0-9]+", "_", str(bundle.get("mine_context", {}).get("mine_name") or "site").strip().lower()).strip("_") or "site"
        base_name = f"mining_maintenance_{site_name}_{report_key}_{timestamp}"
        target = exports_dir() / f"{base_name}.{export_format}"

        try:
            self.status_label.configure(text="Generating report...", text_color="#60a5fa")
            if export_format == "csv":
                self._export_csv(bundle, target)
            elif export_format == "xlsx":
                self._export_excel(bundle, target)
            elif export_format == "html":
                self._export_html(bundle, target)
            elif export_format == "pdf":
                target = self._export_pdf(bundle, target)
            elif export_format == "docx":
                target = self._export_docx(bundle, target)
            self._last_generated_report = Path(target)
            self.status_label.configure(text=f"Report saved to {target}", text_color="#22c55e")
            opened = self._open_report_file(Path(target))
            if opened:
                messagebox.showinfo("Report Generated", f"Report saved and opened:\n{target}")
            else:
                messagebox.showinfo("Report Generated", f"Report saved to:\n{target}")
        except Exception as exc:
            logger.exception("Error generating report: %s", exc)
            self.status_label.configure(text=f"Report generation failed: {exc}", text_color="#ef4444")
            messagebox.showerror("Report Generation Error", f"Failed to generate report:\n{exc}")

    def _open_last_report(self):
        target = self._last_generated_report
        if target is None or not Path(target).exists():
            messagebox.showinfo("Open Last Report", "No generated report available yet. Generate one first.")
            return
        if not self._open_report_file(Path(target)):
            messagebox.showerror("Open Last Report", f"Could not open report:\n{target}")

    def _open_report_file(self, target: Path) -> bool:
        try:
            if not target.exists():
                return False
            if os.name == "nt":
                try:
                    os.startfile(str(target))
                    return True
                except Exception:
                    pass

                ext = target.suffix.lower()
                candidates: List[str] = []
                if ext in (".pdf", ".html", ".htm"):
                    candidates.extend(
                        [
                            r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
                            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
                            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
                        ]
                    )
                elif ext == ".docx":
                    candidates.append(r"C:\Program Files\Microsoft Office\Root\Office16\WINWORD.EXE")
                elif ext in (".xlsx", ".csv"):
                    candidates.append(r"C:\Program Files\Microsoft Office\Root\Office16\EXCEL.EXE")

                for app in candidates:
                    if os.path.isfile(app):
                        try:
                            subprocess.Popen([app, str(target)])
                            return True
                        except Exception:
                            continue
                return False

            opener = "open" if os.uname().sysname.lower() == "darwin" else "xdg-open"
            subprocess.Popen([opener, str(target)])
            return True
        except Exception:
            logger.exception("Failed to open report file: %s", target)
            return False

    def _open_exports_folder(self):
        folder = exports_dir()
        try:
            folder.mkdir(parents=True, exist_ok=True)
            if os.name == "nt":
                os.startfile(str(folder))
            elif os.name == "posix":
                opener = "open" if os.uname().sysname.lower() == "darwin" else "xdg-open"
                subprocess.Popen([opener, str(folder)])
            else:
                raise RuntimeError("Unsupported operating system")
            self.status_label.configure(text=f"Opened exports folder: {folder}", text_color="#60a5fa")
        except Exception as exc:
            logger.exception("Failed to open exports folder: %s", exc)
            messagebox.showerror("Open Exports Folder", f"Failed to open exports folder:\n{exc}")

    def _bundle_rows_for_csv(self, bundle: Dict[str, Any]) -> List[Dict[str, Any]]:
        rows: List[Dict[str, Any]] = []
        for section in bundle.get("sections", []):
            title = section.get("title") or "Section"
            for row in section.get("rows", []) or []:
                item = dict(row)
                item["_section"] = title
                rows.append(item)
        return rows

    def _export_csv(self, bundle: Dict[str, Any], target: Path):
        rows = self._bundle_rows_for_csv(bundle)
        df = pd.DataFrame(rows if rows else [{"_section": bundle.get("title", "Report"), "message": "No rows available"}])
        df.to_csv(target, index=False)

    def _export_excel(self, bundle: Dict[str, Any], target: Path):
        with pd.ExcelWriter(target, engine="openpyxl") as writer:
            pd.DataFrame(bundle.get("summary", [])).to_excel(writer, sheet_name="Summary", index=False)
            for idx, section in enumerate(bundle.get("sections", []), start=1):
                rows = section.get("rows", []) or [{"message": "No rows available"}]
                name = str(section.get("title") or f"Section {idx}")[:31]
                pd.DataFrame(rows).to_excel(writer, sheet_name=name, index=False)

    def _build_html_document(self, bundle: Dict[str, Any]) -> str:
        generated_at = bundle["generated_at"].strftime("%Y-%m-%d %H:%M:%S")
        mine = bundle.get("mine_context", {}) or {}
        mine_name = str(mine.get("mine_name") or mine.get("display_name") or "Unassigned Mine")
        company_name = str(mine.get("company_name") or "").strip()
        quarry_type = str(mine.get("quarry_type") or "").strip()
        lease_area = str(mine.get("lease_area") or "").strip()
        address = str(mine.get("address") or "").strip()
        maps_link = str(mine.get("google_maps_link") or "").strip()
        summary_cards = "\n".join(
            f"""
            <div class="summary-card">
              <div class="summary-label">{escape(str(item.get('label') or ''))}</div>
              <div class="summary-value">{escape(str(item.get('value') or ''))}</div>
            </div>
            """
            for item in bundle.get("summary", [])
        )

        sections_html = []
        for section in bundle.get("sections", []):
            rows = section.get("rows", []) or []
            sections_html.append(
                f"""
                <section class="section">
                  <div class="section-head">
                    <h2>{escape(str(section.get('title') or 'Section'))}</h2>
                    <p>{escape(str(section.get('description') or ''))}</p>
                  </div>
                  {self._rows_to_html_table(rows)}
                </section>
                """
            )

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{escape(str(bundle.get("title") or "Report"))}</title>
  <style>
    :root {{
      --bg: #e2e8f0;
      --paper: #ffffff;
      --ink: #0f172a;
      --muted: #475569;
      --line: #dbe2ea;
      --soft: #f8fafc;
    }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; font-family: "Segoe UI", Arial, sans-serif; background: linear-gradient(180deg, #dbeafe 0%, var(--bg) 60%); color: var(--ink); }}
    .page {{ max-width: 1180px; margin: 28px auto; background: var(--paper); border-radius: 24px; box-shadow: 0 28px 80px rgba(15, 23, 42, 0.14); overflow: hidden; }}
    .hero {{ padding: 30px 38px; background: linear-gradient(135deg, #0f172a, #1d4ed8); color: white; }}
    .eyebrow {{ font-size: 12px; letter-spacing: 0.08em; text-transform: uppercase; opacity: 0.72; }}
    h1 {{ margin: 8px 0 10px; font-size: 32px; }}
    .subtitle {{ margin: 0; max-width: 760px; line-height: 1.6; color: rgba(255,255,255,0.86); }}
    .stamp {{ margin-top: 16px; font-size: 13px; color: rgba(255,255,255,0.78); }}
    .site-bar {{ margin-top: 18px; display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 10px; }}
    .site-chip {{ padding: 10px 12px; border: 1px solid rgba(255,255,255,0.18); border-radius: 14px; background: rgba(255,255,255,0.08); overflow-wrap: anywhere; word-break: break-word; }}
    .site-chip strong {{ display: block; font-size: 11px; letter-spacing: 0.06em; text-transform: uppercase; color: rgba(255,255,255,0.72); margin-bottom: 4px; }}
    .summary-grid {{ display: grid; grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 14px; padding: 24px 34px 4px; }}
    .summary-card {{ background: var(--soft); border: 1px solid var(--line); border-radius: 16px; padding: 16px; }}
    .summary-label {{ font-size: 12px; text-transform: uppercase; letter-spacing: 0.06em; color: var(--muted); margin-bottom: 8px; }}
    .summary-value {{ font-size: 22px; font-weight: 700; overflow-wrap: anywhere; word-break: break-word; }}
    .section {{ padding: 20px 34px 8px; }}
    .section-head h2 {{ margin: 0 0 6px; font-size: 20px; }}
    .section-head p {{ margin: 0 0 14px; color: var(--muted); line-height: 1.6; }}
    table {{ width: 100%; border-collapse: collapse; font-size: 13px; margin-bottom: 12px; table-layout: fixed; }}
    th, td {{ border: 1px solid var(--line); padding: 10px 12px; vertical-align: top; text-align: left; overflow-wrap: anywhere; word-break: break-word; }}
    th {{ background: #eff6ff; color: #1e3a8a; font-weight: 700; }}
    tr:nth-child(even) td {{ background: #f8fafc; }}
    .empty {{ padding: 16px; border: 1px dashed var(--line); border-radius: 12px; color: var(--muted); background: var(--soft); }}
    .footer {{ padding: 18px 34px 30px; font-size: 12px; color: var(--muted); }}
    @media print {{
      body {{ background: white; }}
      .page {{ margin: 0; max-width: none; border-radius: 0; box-shadow: none; }}
      .section {{ page-break-inside: avoid; }}
    }}
    @media (max-width: 900px) {{
      .site-bar {{ grid-template-columns: 1fr; }}
      .summary-grid {{ grid-template-columns: repeat(2, minmax(0, 1fr)); }}
    }}
  </style>
</head>
<body>
  <div class="page">
    <header class="hero">
      <div class="eyebrow">Mining Maintenance System</div>
      <h1>{escape(str(bundle.get("title") or "Report"))}</h1>
      <p class="subtitle">{escape(str(bundle.get("subtitle") or ""))}</p>
      <div class="stamp">Generated on {escape(generated_at)} from user-managed application data.</div>
      <div class="site-bar">
        <div class="site-chip"><strong>Active Mine</strong>{escape(mine_name)}</div>
        <div class="site-chip"><strong>Company</strong>{escape(company_name or "Not set")}</div>
        <div class="site-chip"><strong>Quarry Type</strong>{escape(quarry_type or "Not set")}</div>
        <div class="site-chip"><strong>Lease Area</strong>{escape(lease_area or "Not set")}</div>
        <div class="site-chip"><strong>Address</strong>{escape(address or "Not set")}</div>
        <div class="site-chip"><strong>Google Maps</strong>{escape(maps_link or "Not set")}</div>
      </div>
    </header>
    <section class="summary-grid">
      {summary_cards}
    </section>
    {''.join(sections_html)}
    <div class="footer">This report reflects the same data saved inside the app for machines, schedules, operators, hour entries, and maintenance history.</div>
  </div>
</body>
</html>
"""

    def _rows_to_html_table(self, rows: List[Dict[str, Any]]) -> str:
        if not rows:
            return '<div class="empty">No records available for this section.</div>'
        # Build a stable, complete column list across all rows to avoid
        # dropped fields when later rows contain additional keys.
        columns: List[str] = []
        seen = set()
        for row in rows:
            for key in row.keys():
                if key not in seen:
                    seen.add(key)
                    columns.append(key)
        head = "".join(f"<th>{escape(str(col))}</th>" for col in columns)
        body_rows = []
        for row in rows:
            body_rows.append("<tr>" + "".join(f"<td>{escape(str(row.get(col, '')))}</td>" for col in columns) + "</tr>")
        return f"<table><thead><tr>{head}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"

    def _export_html(self, bundle: Dict[str, Any], target: Path):
        target.write_text(self._build_html_document(bundle), encoding="utf-8")

    def _export_pdf(self, bundle: Dict[str, Any], target: Path) -> Path:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import landscape, letter
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError:
            fallback = target.with_suffix(".html")
            self._export_html(bundle, fallback)
            return fallback

        doc = SimpleDocTemplate(str(target), pagesize=landscape(letter), leftMargin=28, rightMargin=28, topMargin=28, bottomMargin=28)
        styles = getSampleStyleSheet()
        title_style = styles["Heading1"]
        subtitle_style = ParagraphStyle("Subtitle", parent=styles["BodyText"], fontSize=10, textColor=colors.HexColor("#475569"), leading=14)
        story = [
            Paragraph(escape(str(bundle.get("title") or "Report")), title_style),
            Spacer(1, 6),
            Paragraph(escape(str(bundle.get("subtitle") or "")), subtitle_style),
            Spacer(1, 12),
        ]
        mine = bundle.get("mine_context", {}) or {}
        site_rows = [
            ["Active Mine", str(mine.get("mine_name") or mine.get("display_name") or "Unassigned Mine")],
            ["Company", str(mine.get("company_name") or "Not set")],
            ["Quarry Type", str(mine.get("quarry_type") or "Not set")],
            ["Lease Area", str(mine.get("lease_area") or "Not set")],
            ["Address", str(mine.get("address") or "Not set")],
            ["Google Maps", str(mine.get("google_maps_link") or "Not set")],
        ]
        site_table = Table([["Site Context", "Details"]] + site_rows, repeatRows=1)
        site_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F8FAFC"), colors.white]),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
        ]))
        story.extend([site_table, Spacer(1, 14)])

        summary_rows = [["Metric", "Value"]] + [[str(item.get("label") or ""), str(item.get("value") or "")] for item in bundle.get("summary", [])]
        summary_table = Table(summary_rows, repeatRows=1)
        summary_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DBEAFE")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
        ]))
        story.extend([summary_table, Spacer(1, 14)])

        for section in bundle.get("sections", []):
            story.append(Paragraph(escape(str(section.get("title") or "Section")), styles["Heading2"]))
            if section.get("description"):
                story.append(Paragraph(escape(str(section.get("description") or "")), subtitle_style))
            rows = section.get("rows", []) or []
            if not rows:
                story.extend([Paragraph("No records available for this section.", styles["BodyText"]), Spacer(1, 10)])
                continue
            columns = list(rows[0].keys())
            table_rows = [columns] + [[str(row.get(col, "")) for col in columns] for row in rows]
            table = Table(table_rows, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.extend([Spacer(1, 6), table, Spacer(1, 12)])

        doc.build(story)
        return target

    def _export_docx(self, bundle: Dict[str, Any], target: Path) -> Path:
        try:
            from docx import Document
        except ImportError:
            fallback = target.with_suffix(".html")
            self._export_html(bundle, fallback)
            return fallback

        doc = Document()
        doc.add_heading(str(bundle.get("title") or "Report"), 0)
        doc.add_paragraph(str(bundle.get("subtitle") or ""))
        doc.add_paragraph(f"Generated on: {bundle['generated_at'].strftime('%Y-%m-%d %H:%M:%S')}")
        mine = bundle.get("mine_context", {}) or {}
        doc.add_heading("Mine Site Context", level=1)
        mine_table = doc.add_table(rows=1, cols=2)
        mine_table.style = "Table Grid"
        mine_hdr = mine_table.rows[0].cells
        mine_hdr[0].text = "Field"
        mine_hdr[1].text = "Value"
        for label, value in (
            ("Active Mine", str(mine.get("mine_name") or mine.get("display_name") or "Unassigned Mine")),
            ("Company", str(mine.get("company_name") or "Not set")),
            ("Quarry Type", str(mine.get("quarry_type") or "Not set")),
            ("Lease Area", str(mine.get("lease_area") or "Not set")),
            ("Address", str(mine.get("address") or "Not set")),
            ("Google Maps", str(mine.get("google_maps_link") or "Not set")),
        ):
            row = mine_table.add_row().cells
            row[0].text = label
            row[1].text = value

        doc.add_heading("Summary", level=1)
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = "Table Grid"
        hdr = summary_table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"
        for item in bundle.get("summary", []):
            row = summary_table.add_row().cells
            row[0].text = str(item.get("label") or "")
            row[1].text = str(item.get("value") or "")

        for section in bundle.get("sections", []):
            doc.add_heading(str(section.get("title") or "Section"), level=1)
            if section.get("description"):
                doc.add_paragraph(str(section.get("description") or ""))
            rows = section.get("rows", []) or []
            if not rows:
                doc.add_paragraph("No records available for this section.")
                continue
            columns = list(rows[0].keys())
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = "Table Grid"
            for idx, col in enumerate(columns):
                table.rows[0].cells[idx].text = str(col)
            for item in rows:
                cells = table.add_row().cells
                for idx, col in enumerate(columns):
                    cells[idx].text = str(item.get(col, ""))

        doc.save(target)
        return target
