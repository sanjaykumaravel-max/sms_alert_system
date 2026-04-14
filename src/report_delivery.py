from __future__ import annotations

import json
import logging
import os
import smtplib
from datetime import datetime
from email.message import EmailMessage
from html import escape
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from .app_paths import data_path, exports_dir
    from .machine_store import evaluate_machine_status, load_machines
except Exception:
    from app_paths import data_path, exports_dir
    from machine_store import evaluate_machine_status, load_machines


logger = logging.getLogger(__name__)
REPORT_STATE_FILE = data_path("report_delivery_state.json")


def _now() -> datetime:
    return datetime.now()


def _parse_csv_items(value: Any) -> List[str]:
    parts = str(value or "").replace(";", ",").replace("\n", ",").split(",")
    return [item.strip() for item in parts if item and item.strip()]


def _report_slot_key(*, now: datetime, frequency: str, weekday: int, hour: int) -> Optional[str]:
    if now.hour < hour:
        return None
    if frequency == "weekly":
        if now.weekday() != weekday:
            return None
        year, week, _ = now.isocalendar()
        return f"weekly:{year}-W{week:02d}"
    return f"daily:{now.date().isoformat()}"


def _load_delivery_state(path: Path = REPORT_STATE_FILE) -> Dict[str, Any]:
    try:
        if path.exists():
            payload = json.loads(path.read_text(encoding="utf-8")) or {}
            if isinstance(payload, dict):
                return payload
    except Exception:
        logger.exception("Failed to load report delivery state")
    return {}


def _save_delivery_state(state: Dict[str, Any], path: Path = REPORT_STATE_FILE) -> None:
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(state, indent=2), encoding="utf-8")
    except Exception:
        logger.exception("Failed to save report delivery state")


def _maintenance_rows() -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for machine in load_machines():
        ctx = evaluate_machine_status(machine)
        status = str(ctx.get("status") or "normal").lower()
        rows.append(
            {
                "machine_id": machine.get("id") or "",
                "machine_name": machine.get("name") or machine.get("model") or machine.get("id") or "",
                "status": status,
                "trigger": str(ctx.get("trigger") or "manual").lower(),
                "due_date": ctx.get("due_date") or machine.get("due_date") or machine.get("next_maintenance") or "",
                "current_hours": ctx.get("current_hours") or machine.get("current_hours") or machine.get("hours") or "",
                "next_due_hours": ctx.get("next_due_hours") or machine.get("next_due_hours") or "",
                "operator_phone": machine.get("operator_phone") or "",
            }
        )
    rank = {"critical": 0, "overdue": 1, "due": 2, "maintenance": 3, "normal": 4}
    rows.sort(key=lambda row: (rank.get(str(row.get("status") or "normal"), 9), str(row.get("machine_id") or "")))
    return rows


def _build_html(rows: List[Dict[str, Any]], generated_at: datetime) -> str:
    total = len(rows)
    active = [row for row in rows if str(row.get("status") or "normal") in {"maintenance", "due", "overdue", "critical"}]
    overdue = [row for row in rows if str(row.get("status") or "normal") == "overdue"]
    headers = ("Machine ID", "Machine", "Status", "Trigger", "Due Date", "Current Hours", "Next Due Hours", "Operator Phone")
    body = []
    for row in rows:
        body.append(
            "<tr>"
            f"<td>{escape(str(row.get('machine_id') or ''))}</td>"
            f"<td>{escape(str(row.get('machine_name') or ''))}</td>"
            f"<td>{escape(str(row.get('status') or '').title())}</td>"
            f"<td>{escape(str(row.get('trigger') or '').title())}</td>"
            f"<td>{escape(str(row.get('due_date') or ''))}</td>"
            f"<td>{escape(str(row.get('current_hours') or ''))}</td>"
            f"<td>{escape(str(row.get('next_due_hours') or ''))}</td>"
            f"<td>{escape(str(row.get('operator_phone') or ''))}</td>"
            "</tr>"
        )
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Maintenance Report</title>
  <style>
    body {{ font-family: "Segoe UI", Arial, sans-serif; margin: 0; background: #f1f5f9; color: #0f172a; }}
    .page {{ max-width: 1180px; margin: 24px auto; background: #fff; border-radius: 18px; box-shadow: 0 20px 50px rgba(15, 23, 42, 0.12); overflow: hidden; }}
    .hero {{ padding: 24px 30px; background: linear-gradient(135deg, #0f172a, #1d4ed8); color: #fff; }}
    .hero h1 {{ margin: 0 0 8px; font-size: 28px; }}
    .hero p {{ margin: 0; opacity: 0.9; }}
    .summary {{ display: grid; grid-template-columns: repeat(3, minmax(0, 1fr)); gap: 12px; padding: 18px 22px; }}
    .tile {{ background: #f8fafc; border: 1px solid #dbe2ea; border-radius: 12px; padding: 12px; }}
    .tile .label {{ color: #475569; font-size: 12px; text-transform: uppercase; letter-spacing: 0.06em; }}
    .tile .value {{ font-size: 22px; font-weight: 700; margin-top: 5px; }}
    .wrap {{ padding: 0 22px 22px; }}
    table {{ width: 100%; border-collapse: collapse; table-layout: fixed; }}
    th, td {{ border: 1px solid #dbe2ea; padding: 9px 10px; text-align: left; vertical-align: top; overflow-wrap: anywhere; }}
    th {{ background: #eff6ff; color: #1e3a8a; font-size: 12px; text-transform: uppercase; letter-spacing: 0.04em; }}
    tr:nth-child(even) td {{ background: #f8fafc; }}
  </style>
</head>
<body>
  <div class="page">
    <header class="hero">
      <h1>Automatic Maintenance Report</h1>
      <p>Generated at {escape(generated_at.strftime("%Y-%m-%d %H:%M:%S"))}</p>
    </header>
    <section class="summary">
      <div class="tile"><div class="label">Machines</div><div class="value">{total}</div></div>
      <div class="tile"><div class="label">Maintenance Active</div><div class="value">{len(active)}</div></div>
      <div class="tile"><div class="label">Overdue</div><div class="value">{len(overdue)}</div></div>
    </section>
    <div class="wrap">
      <table>
        <thead><tr>{''.join(f'<th>{escape(h)}</th>' for h in headers)}</tr></thead>
        <tbody>{''.join(body) or '<tr><td colspan="8">No machine data found.</td></tr>'}</tbody>
      </table>
    </div>
  </div>
</body>
</html>"""


def _export_report(*, rows: List[Dict[str, Any]], export_format: str, generated_at: datetime) -> Path:
    stamp = generated_at.strftime("%Y%m%d_%H%M%S")
    exports = exports_dir()
    exports.mkdir(parents=True, exist_ok=True)

    fmt = str(export_format or "pdf").strip().lower()
    target = exports / f"auto_maintenance_report_{stamp}.{fmt}"
    html_doc = _build_html(rows, generated_at)

    if fmt in {"html", "htm"}:
        target.write_text(html_doc, encoding="utf-8")
        return target

    if fmt == "pdf":
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import landscape, letter
            from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
        except Exception:
            fallback = target.with_suffix(".html")
            fallback.write_text(html_doc, encoding="utf-8")
            return fallback

        doc = SimpleDocTemplate(str(target), pagesize=landscape(letter), leftMargin=24, rightMargin=24, topMargin=24, bottomMargin=24)
        styles = getSampleStyleSheet()
        story = [
            Paragraph("Automatic Maintenance Report", styles["Heading1"]),
            Paragraph(f"Generated at {generated_at.strftime('%Y-%m-%d %H:%M:%S')}", styles["BodyText"]),
            Spacer(1, 10),
        ]
        columns = ["machine_id", "machine_name", "status", "trigger", "due_date", "current_hours", "next_due_hours", "operator_phone"]
        header = ["Machine ID", "Machine", "Status", "Trigger", "Due Date", "Current Hours", "Next Due Hours", "Operator Phone"]
        table_rows = [header]
        for row in rows:
            table_rows.append([str(row.get(col, "")) for col in columns])
        if len(table_rows) == 1:
            table_rows.append(["No machine data found.", "", "", "", "", "", "", ""])
        table = Table(table_rows, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(table)
        doc.build(story)
        return target

    if fmt == "csv":
        import csv

        with target.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(["machine_id", "machine_name", "status", "trigger", "due_date", "current_hours", "next_due_hours", "operator_phone"])
            for row in rows:
                writer.writerow(
                    [
                        row.get("machine_id", ""),
                        row.get("machine_name", ""),
                        row.get("status", ""),
                        row.get("trigger", ""),
                        row.get("due_date", ""),
                        row.get("current_hours", ""),
                        row.get("next_due_hours", ""),
                        row.get("operator_phone", ""),
                    ]
                )
        return target

    if fmt == "xlsx":
        import pandas as pd

        pd.DataFrame(rows).to_excel(target, index=False)
        return target

    if fmt == "docx":
        try:
            from docx import Document
        except Exception:
            fallback = target.with_suffix(".html")
            fallback.write_text(html_doc, encoding="utf-8")
            return fallback
        doc = Document()
        doc.add_heading("Automatic Maintenance Report", 0)
        doc.add_paragraph(f"Generated at {generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        headers = ["Machine ID", "Machine", "Status", "Trigger", "Due Date", "Current Hours", "Next Due Hours", "Operator Phone"]
        for idx, item in enumerate(headers):
            table.rows[0].cells[idx].text = item
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row.get("machine_id", ""))
            cells[1].text = str(row.get("machine_name", ""))
            cells[2].text = str(row.get("status", "")).title()
            cells[3].text = str(row.get("trigger", "")).title()
            cells[4].text = str(row.get("due_date", ""))
            cells[5].text = str(row.get("current_hours", ""))
            cells[6].text = str(row.get("next_due_hours", ""))
            cells[7].text = str(row.get("operator_phone", ""))
        doc.save(str(target))
        return target

    fallback = target.with_suffix(".html")
    fallback.write_text(html_doc, encoding="utf-8")
    return fallback


def _smtp_config(settings: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "host": str(settings.get("smtp_host") or os.getenv("SMTP_HOST") or "").strip(),
        "port": int(settings.get("smtp_port") or os.getenv("SMTP_PORT") or 587),
        "username": str(settings.get("smtp_username") or os.getenv("SMTP_USERNAME") or "").strip(),
        "password": str(settings.get("smtp_password") or os.getenv("SMTP_PASSWORD") or "").strip(),
        "sender": str(settings.get("smtp_sender_email") or os.getenv("SMTP_SENDER_EMAIL") or "").strip(),
        "use_tls": bool(settings.get("smtp_use_tls", True)),
    }


def _send_report_email(
    *,
    recipients: List[str],
    subject: str,
    body: str,
    attachment: Path,
    smtp_cfg: Dict[str, Any],
) -> Dict[str, Any]:
    if not recipients:
        return {"success": False, "reason": "no_recipients"}
    host = str(smtp_cfg.get("host") or "").strip()
    sender = str(smtp_cfg.get("sender") or smtp_cfg.get("username") or "").strip()
    if not host or not sender:
        return {"success": False, "reason": "smtp_not_configured"}

    msg = EmailMessage()
    msg["From"] = sender
    msg["To"] = ", ".join(recipients)
    msg["Subject"] = subject or "Automatic Maintenance Report"
    msg.set_content(body or "Please find the attached maintenance report.")
    msg.add_attachment(
        attachment.read_bytes(),
        maintype="application",
        subtype="octet-stream",
        filename=attachment.name,
    )

    port = int(smtp_cfg.get("port") or 587)
    username = str(smtp_cfg.get("username") or "").strip()
    password = str(smtp_cfg.get("password") or "").strip()
    use_tls = bool(smtp_cfg.get("use_tls", True))

    server = smtplib.SMTP(host=host, port=port, timeout=30)
    try:
        server.ehlo()
        if use_tls:
            server.starttls()
            server.ehlo()
        if username:
            server.login(username, password)
        server.send_message(msg)
    finally:
        try:
            server.quit()
        except Exception:
            pass
    return {"success": True, "reason": "sent"}


def maybe_deliver_scheduled_report(
    *,
    settings: Optional[Dict[str, Any]] = None,
    now: Optional[datetime] = None,
    state_file: Path = REPORT_STATE_FILE,
) -> Dict[str, Any]:
    cfg = dict(settings or {})
    summary: Dict[str, Any] = {
        "triggered": False,
        "scheduled": False,
        "reason": "",
        "report_file": None,
        "email_sent": False,
    }

    if not bool(cfg.get("auto_report_delivery_enabled", False)):
        summary["reason"] = "report_delivery_disabled"
        return summary

    current = now or _now()
    frequency = str(cfg.get("report_delivery_frequency") or "daily").strip().lower()
    if frequency not in {"daily", "weekly"}:
        frequency = "daily"
    hour = int(cfg.get("report_delivery_hour", 18) or 18)
    weekday = int(cfg.get("report_delivery_weekday", 0) or 0)
    slot_key = _report_slot_key(now=current, frequency=frequency, weekday=weekday, hour=hour)
    if not slot_key:
        summary["reason"] = "slot_not_due"
        return summary

    state = _load_delivery_state(state_file)
    if str(state.get("last_slot_key") or "") == slot_key:
        summary["reason"] = "already_sent_for_slot"
        return summary

    rows = _maintenance_rows()
    report_format = str(cfg.get("report_delivery_format") or "pdf").strip().lower()
    try:
        report_file = _export_report(rows=rows, export_format=report_format, generated_at=current)
    except Exception as exc:
        logger.exception("Scheduled report export failed")
        state.update(
            {
                "last_slot_key": slot_key,
                "last_run_time": current.isoformat(timespec="seconds"),
                "last_result": f"export_failed:{exc}",
                "last_email_sent": False,
                "last_email_reason": f"export_failed:{exc}",
            }
        )
        _save_delivery_state(state, state_file)
        summary["triggered"] = True
        summary["scheduled"] = False
        summary["reason"] = f"export_failed:{exc}"
        return summary

    recipients = _parse_csv_items(cfg.get("report_delivery_emails"))
    subject = str(cfg.get("report_delivery_email_subject") or "Mining Maintenance Report").strip()
    body = (
        f"Automatic {frequency} maintenance report generated at {current.strftime('%Y-%m-%d %H:%M:%S')}.\n"
        f"Machines: {len(rows)}\n"
        f"Attachment: {report_file.name}"
    )
    smtp_cfg = _smtp_config(cfg)
    try:
        email_result = _send_report_email(
            recipients=recipients,
            subject=subject,
            body=body,
            attachment=report_file,
            smtp_cfg=smtp_cfg,
        )
    except Exception as exc:
        logger.exception("Scheduled report email send failed")
        email_result = {"success": False, "reason": f"email_error:{exc}"}

    state.update(
        {
            "last_slot_key": slot_key,
            "last_run_time": current.isoformat(timespec="seconds"),
            "last_report_file": str(report_file),
            "last_result": "sent" if email_result.get("success") else str(email_result.get("reason") or "generated_only"),
            "last_email_sent": bool(email_result.get("success")),
            "last_email_reason": str(email_result.get("reason") or ""),
        }
    )
    _save_delivery_state(state, state_file)

    summary["triggered"] = True
    summary["scheduled"] = True
    summary["reason"] = str(email_result.get("reason") or "generated")
    summary["report_file"] = str(report_file)
    summary["email_sent"] = bool(email_result.get("success"))
    return summary
