from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
SOURCE = DOCS_DIR / "PROJECT_REVIEW_REPORT_6_CHAPTERS.md"
DOCX_OUT = DOCS_DIR / "PROJECT_REVIEW_REPORT_6_CHAPTERS.docx"
HTML_OUT = DOCS_DIR / "PROJECT_REVIEW_REPORT_6_CHAPTERS.html"


def _inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def build_docx(markdown_text: str) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_done = False
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            title_done = True
            continue

        if stripped.startswith("## "):
            level = 1 if title_done else 0
            p = doc.add_paragraph()
            p.style = f"Heading {level}"
            _inline_runs(p, stripped[3:].strip())
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.style = "Heading 2"
            _inline_runs(p, stripped[4:].strip())
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _inline_runs(p, stripped[2:].strip())
            continue

        p = doc.add_paragraph()
        _inline_runs(p, stripped.replace("  ", " "))

    doc.save(DOCX_OUT)


def build_html(markdown_text: str) -> None:
    body: list[str] = []
    in_list = False
    ordered_list = False

    def close_list() -> None:
        nonlocal in_list, ordered_list
        if in_list:
            body.append("</ol>" if ordered_list else "</ul>")
            in_list = False
            ordered_list = False

    def render_inline(text: str) -> str:
        escaped = html.escape(text)
        return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)

    for raw_line in markdown_text.splitlines():
        stripped = raw_line.strip()

        if not stripped or stripped == "---":
            close_list()
            continue

        if stripped.startswith("# "):
            close_list()
            body.append(f"<h1>{render_inline(stripped[2:].strip())}</h1>")
            continue

        if stripped.startswith("## "):
            close_list()
            body.append(f"<h2>{render_inline(stripped[3:].strip())}</h2>")
            continue

        if stripped.startswith("### "):
            close_list()
            body.append(f"<h3>{render_inline(stripped[4:].strip())}</h3>")
            continue

        if re.match(r"^\d+\.\s+", stripped):
            item = re.sub(r"^\d+\.\s+", "", stripped)
            if not in_list or not ordered_list:
                close_list()
                body.append("<ol>")
                in_list = True
                ordered_list = True
            body.append(f"<li>{render_inline(item)}</li>")
            continue

        if stripped.startswith("- "):
            if not in_list or ordered_list:
                close_list()
                body.append("<ul>")
                in_list = True
                ordered_list = False
            body.append(f"<li>{render_inline(stripped[2:].strip())}</li>")
            continue

        close_list()
        body.append(f"<p>{render_inline(stripped)}</p>")

    close_list()

    html_doc = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Project Review Report</title>
  <style>
    :root {{
      --bg: #eef4fb;
      --paper: #ffffff;
      --ink: #102033;
      --muted: #53657d;
      --line: #d8e2ee;
      --accent: #0f4c81;
    }}
    body {{
      margin: 0;
      background: linear-gradient(180deg, #e8f1fb 0%, #f5f8fc 100%);
      color: var(--ink);
      font-family: "Cambria", "Georgia", serif;
      line-height: 1.65;
    }}
    .page {{
      max-width: 920px;
      margin: 28px auto;
      background: var(--paper);
      border: 1px solid var(--line);
      box-shadow: 0 18px 40px rgba(16, 32, 51, 0.08);
      padding: 40px 54px;
    }}
    h1, h2, h3 {{
      color: var(--accent);
      font-family: "Calibri", "Segoe UI", sans-serif;
    }}
    h1 {{
      text-align: center;
      font-size: 28px;
      margin-bottom: 6px;
    }}
    h2 {{
      font-size: 22px;
      margin-top: 28px;
      border-bottom: 1px solid var(--line);
      padding-bottom: 6px;
    }}
    h3 {{
      font-size: 17px;
      margin-top: 20px;
      margin-bottom: 6px;
    }}
    p {{
      margin: 10px 0;
      text-align: justify;
    }}
    ul, ol {{
      margin: 8px 0 14px 24px;
    }}
    li {{
      margin: 4px 0;
    }}
    strong {{
      color: #0c3356;
    }}
    @media print {{
      body {{
        background: #fff;
      }}
      .page {{
        margin: 0;
        box-shadow: none;
        border: none;
        max-width: none;
      }}
    }}
  </style>
</head>
<body>
  <main class="page">
    {''.join(body)}
  </main>
</body>
</html>
"""
    HTML_OUT.write_text(html_doc, encoding="utf-8")


def main() -> None:
    markdown_text = SOURCE.read_text(encoding="utf-8")
    build_docx(markdown_text)
    build_html(markdown_text)
    print(DOCX_OUT)
    print(HTML_OUT)


if __name__ == "__main__":
    main()
